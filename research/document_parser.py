import os
import re
import json
import logging
from typing import Dict, Any, List, Optional
from abc import ABC, abstractmethod

logger = logging.getLogger(__name__)


class DocumentParser(ABC):
    SUPPORTED_EXTENSIONS = []
    
    @abstractmethod
    def parse(self, file_path: str) -> Dict[str, Any]:
        pass
    
    @abstractmethod
    def extract_text(self, file_path: str) -> str:
        pass
    
    @abstractmethod
    def extract_tables(self, file_path: str) -> List[Dict[str, Any]]:
        pass
    
    @abstractmethod
    def extract_images(self, file_path: str) -> List[Dict[str, Any]]:
        pass


class PDFParser(DocumentParser):
    SUPPORTED_EXTENSIONS = ['.pdf']
    
    def __init__(self):
        try:
            import fitz
            self.fitz = fitz
            self.use_fitz = True
            logger.info("Using PyMuPDF for PDF parsing")
        except ImportError:
            try:
                from pdfplumber import open as pdf_open
                self.pdf_open = pdf_open
                self.use_fitz = False
                logger.info("Using pdfplumber for PDF parsing")
            except ImportError:
                self.use_fitz = False
                logger.warning("No PDF parsing library available")
    
    def parse(self, file_path: str) -> Dict[str, Any]:
        if not os.path.exists(file_path):
            return {'error': f"File not found: {file_path}"}
        
        try:
            return {
                'file_path': file_path,
                'file_name': os.path.basename(file_path),
                'text': self.extract_text(file_path),
                'tables': self.extract_tables(file_path),
                'images': self.extract_images(file_path),
                'metadata': self.extract_metadata(file_path),
            }
        except Exception as e:
            logger.error(f"Error parsing PDF {file_path}: {str(e)}")
            return {'error': str(e)}
    
    def extract_text(self, file_path: str) -> str:
        if self.use_fitz:
            return self._extract_text_fitz(file_path)
        else:
            return self._extract_text_pdfplumber(file_path)
    
    def _extract_text_fitz(self, file_path: str) -> str:
        doc = self.fitz.open(file_path)
        text = ""
        for page in doc:
            text += page.get_text() + "\n"
        doc.close()
        return text
    
    def _extract_text_pdfplumber(self, file_path: str) -> str:
        try:
            with self.pdf_open(file_path) as pdf:
                text = ""
                for page in pdf.pages:
                    text += page.extract_text() + "\n"
            return text
        except Exception:
            return ""
    
    def extract_tables(self, file_path: str) -> List[Dict[str, Any]]:
        tables = []
        
        if self.use_fitz:
            doc = self.fitz.open(file_path)
            for page_num, page in enumerate(doc):
                page_tables = page.find_tables()
                for table_idx, table in enumerate(page_tables):
                    try:
                        table_data = []
                        for row in table.extract():
                            table_data.append([str(cell) for cell in row if cell is not None])
                        
                        if table_data and len(table_data) > 1:
                            tables.append({
                                'page': page_num + 1,
                                'table_index': table_idx,
                                'rows': len(table_data),
                                'columns': len(table_data[0]) if table_data else 0,
                                'data': table_data,
                            })
                    except Exception as e:
                        logger.debug(f"Error extracting table {table_idx} on page {page_num + 1}: {str(e)}")
            doc.close()
        else:
            try:
                with self.pdf_open(file_path) as pdf:
                    for page_num, page in enumerate(pdf.pages):
                        page_tables = page.extract_tables()
                        for table_idx, table_data in enumerate(page_tables):
                            if table_data and len(table_data) > 1:
                                tables.append({
                                    'page': page_num + 1,
                                    'table_index': table_idx,
                                    'rows': len(table_data),
                                    'columns': len(table_data[0]) if table_data else 0,
                                    'data': [[str(cell) if cell else '' for cell in row] for row in table_data],
                                })
            except Exception as e:
                logger.debug(f"Error extracting tables with pdfplumber: {str(e)}")
        
        return tables
    
    def extract_images(self, file_path: str) -> List[Dict[str, Any]]:
        images = []
        
        if self.use_fitz:
            doc = self.fitz.open(file_path)
            for page_num, page in enumerate(doc):
                page_images = page.get_images(full=True)
                for img_idx, img in enumerate(page_images):
                    xref = img[0]
                    base_image = doc.extract_image(xref)
                    if base_image:
                        images.append({
                            'page': page_num + 1,
                            'image_index': img_idx,
                            'width': base_image.get('width', 0),
                            'height': base_image.get('height', 0),
                            'size': len(base_image.get('image', b'')),
                            'ext': base_image.get('ext', 'png'),
                        })
            doc.close()
        
        return images
    
    def extract_metadata(self, file_path: str) -> Dict[str, Any]:
        metadata = {}
        
        if self.use_fitz:
            doc = self.fitz.open(file_path)
            info = doc.metadata
            metadata = {
                'title': info.get('title', ''),
                'author': info.get('author', ''),
                'subject': info.get('subject', ''),
                'keywords': info.get('keywords', ''),
                'creator': info.get('creator', ''),
                'producer': info.get('producer', ''),
                'creation_date': info.get('creationDate', ''),
                'modification_date': info.get('modDate', ''),
                'page_count': doc.page_count,
            }
            doc.close()
        
        return metadata


class DOCXParser(DocumentParser):
    SUPPORTED_EXTENSIONS = ['.docx']
    
    def __init__(self):
        try:
            from docx import Document
            self.Document = Document
            self.available = True
            logger.info("Using python-docx for DOCX parsing")
        except ImportError:
            self.available = False
            logger.warning("python-docx not available")
    
    def parse(self, file_path: str) -> Dict[str, Any]:
        if not self.available:
            return {'error': 'python-docx not available'}
        
        if not os.path.exists(file_path):
            return {'error': f"File not found: {file_path}"}
        
        try:
            doc = self.Document(file_path)
            return {
                'file_path': file_path,
                'file_name': os.path.basename(file_path),
                'text': self.extract_text(file_path),
                'tables': self.extract_tables(file_path),
                'images': self.extract_images(file_path),
                'metadata': self._extract_metadata(doc),
            }
        except Exception as e:
            logger.error(f"Error parsing DOCX {file_path}: {str(e)}")
            return {'error': str(e)}
    
    def extract_text(self, file_path: str) -> str:
        if not self.available:
            return ""
        
        doc = self.Document(file_path)
        text = ""
        
        for element in doc.element.body:
            tag = element.tag.split('}')[-1]
            if tag == 'p':
                text += element.text if element.text else ""
                for child in element:
                    child_tag = child.tag.split('}')[-1]
                    if child_tag == 'r':
                        for text_child in child:
                            if text_child.text:
                                text += text_child.text
                text += "\n"
            elif tag == 'tbl':
                text += "\n[TABLE]\n"
        
        return text
    
    def extract_tables(self, file_path: str) -> List[Dict[str, Any]]:
        if not self.available:
            return []
        
        doc = self.Document(file_path)
        tables = []
        
        for table_idx, table in enumerate(doc.tables):
            table_data = []
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    row_data.append(cell.text.strip())
                table_data.append(row_data)
            
            if table_data and len(table_data) > 1:
                tables.append({
                    'table_index': table_idx,
                    'rows': len(table_data),
                    'columns': len(table_data[0]) if table_data else 0,
                    'data': table_data,
                })
        
        return tables
    
    def extract_images(self, file_path: str) -> List[Dict[str, Any]]:
        if not self.available:
            return []
        
        doc = self.Document(file_path)
        images = []
        
        for rel in doc.part.rels.values():
            if "image" in rel.target_ref:
                images.append({
                    'image_index': len(images),
                    'size': len(rel.target_part.blob),
                    'ext': rel.target_ref.split('.')[-1] if '.' in rel.target_ref else 'png',
                })
        
        return images
    
    def _extract_metadata(self, doc) -> Dict[str, Any]:
        core_props = doc.core_properties
        return {
            'title': core_props.title,
            'author': core_props.author,
            'subject': core_props.subject,
            'keywords': core_props.keywords,
            'created': str(core_props.created) if core_props.created else '',
            'modified': str(core_props.modified) if core_props.modified else '',
        }


class TXTParser(DocumentParser):
    SUPPORTED_EXTENSIONS = ['.txt', '.md']
    
    def parse(self, file_path: str) -> Dict[str, Any]:
        if not os.path.exists(file_path):
            return {'error': f"File not found: {file_path}"}
        
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                text = f.read()
            
            return {
                'file_path': file_path,
                'file_name': os.path.basename(file_path),
                'text': text,
                'tables': self.extract_tables_from_text(text),
                'images': [],
                'metadata': {
                    'file_size': os.path.getsize(file_path),
                    'line_count': text.count('\n'),
                },
            }
        except Exception as e:
            logger.error(f"Error parsing TXT {file_path}: {str(e)}")
            return {'error': str(e)}
    
    def extract_text(self, file_path: str) -> str:
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                return f.read()
        except Exception:
            return ""
    
    def extract_tables(self, file_path: str) -> List[Dict[str, Any]]:
        text = self.extract_text(file_path)
        return self.extract_tables_from_text(text)
    
    def extract_tables_from_text(self, text: str) -> List[Dict[str, Any]]:
        tables = []
        lines = text.split('\n')
        table_pattern = re.compile(r'^[\|+\-]')
        
        current_table = []
        for line in lines:
            if table_pattern.match(line):
                current_table.append(line)
            elif current_table:
                parsed_table = self._parse_text_table(current_table)
                if parsed_table:
                    tables.append(parsed_table)
                current_table = []
        
        if current_table:
            parsed_table = self._parse_text_table(current_table)
            if parsed_table:
                tables.append(parsed_table)
        
        return tables
    
    def _parse_text_table(self, lines: List[str]) -> Optional[Dict[str, Any]]:
        table_data = []
        for line in lines:
            cleaned = line.strip().strip('|')
            if cleaned and not cleaned.startswith('+') and not cleaned.startswith('-'):
                cells = [cell.strip() for cell in cleaned.split('|')]
                table_data.append(cells)
        
        if table_data and len(table_data) > 1:
            return {
                'rows': len(table_data),
                'columns': len(table_data[0]) if table_data else 0,
                'data': table_data,
            }
        return None
    
    def extract_images(self, file_path: str) -> List[Dict[str, Any]]:
        return []


class DocumentParserFactory:
    _parsers = {}
    
    @classmethod
    def register_parser(cls, parser_class: type):
        for ext in parser_class.SUPPORTED_EXTENSIONS:
            cls._parsers[ext.lower()] = parser_class
    
    @classmethod
    def get_parser(cls, file_path: str) -> Optional[DocumentParser]:
        _, ext = os.path.splitext(file_path)
        parser_class = cls._parsers.get(ext.lower())
        if parser_class:
            return parser_class()
        return None
    
    @classmethod
    def parse(cls, file_path: str) -> Dict[str, Any]:
        parser = cls.get_parser(file_path)
        if parser:
            return parser.parse(file_path)
        return {'error': f"Unsupported file format: {file_path}"}


DocumentParserFactory.register_parser(PDFParser)
DocumentParserFactory.register_parser(DOCXParser)
DocumentParserFactory.register_parser(TXTParser)